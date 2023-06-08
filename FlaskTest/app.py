from fileinput import filename
import time
from flask import Flask, render_template, request, Response,redirect,send_file
import json
import os

import pandas as pd
import math
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import docx
from docx import Document

import sys
import subprocess
import psutil

def createPDF(i):    
    doc = Document(input_docx)
    print(i)
    #Grab first table
    table = doc.tables[0]
    
    #set Agent Name
    cell = table.cell (2,1)
    cell.text = ": "+dfs.iloc[i]['Agent Name ']
    
    #set Account Holder Name
    cell = table.cell(2,3)
    cell.text = ": "+str(dfs.iloc[i]['Account Name'])
    
    #set Email
    cell = table.cell(3,1)
    cell.text = ": "+ dfs.iloc[i]['Email']
    
    #set Bank Name
    cell = table.cell(3,3)
    cell.text = ": AYA"
    
    #set Position
    cell = table.cell(4,1)
    cell.text = ": "+dfs.iloc[i]['Position']
    
    #set Bank Account No
    cell = table.cell(4,3)
    cell.text = ": "+ str(dfs.iloc[i]['Bank Acc'])
    
    #set NRIC/Passport
    cell = table.cell(5,1)
    cell.text = ": "+ str(dfs.iloc[i]['NRC / Passport'])
    
    #set Agent Code
    cell = table.cell(6,1)
    cell.text = ": "+ dfs.iloc[i]['Agent Code']
    
    #For First Table, description and amount (mmk)
    j = 9
    
    #set FYC if exists
   # if(dfs.iloc[i]['FYC']!=0):
    if(not(pd.isnull(dfs.iloc[i]['FYC']))==True and (dfs.iloc[i]['FYC']!='0')):
        cell = table.cell(j,0)
        cell.text = "FYC"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['FYC'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        j+=1
        
    #set RYC if exists    
    if(not(pd.isnull(dfs.iloc[i]['RYC']))==True and (dfs.iloc[i]['RYC']!='0')):
        cell = table.cell(j,0)
        cell.text = "RYC"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['RYC'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
    
    #set Direct Overriding if exists
    if(not(pd.isnull(dfs.iloc[i]['Direct Overriding']))==True and (dfs.iloc[i]['Direct Overriding']!='0')):
        cell = table.cell(j,0)
        cell.text = "Direct Overriding"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['Direct Overriding'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
        
    #set Indirect Overriding 1 if exists
    if(not(pd.isnull(dfs.iloc[i]['Indirect Overriding 1']))==True and (dfs.iloc[i]['Indirect Overriding 1']!='0')):
        cell = table.cell(j,0)
        cell.text = "Indirect Overriding 1"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['Indirect Overriding 1'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
        
    #set Indirect Overriding 2 if exists
    if(not(pd.isnull(dfs.iloc[i]['Indirect Overriding 2']))==True and (dfs.iloc[i]['Indirect Overriding 2']!='0')):
        cell = table.cell(j,0)
        cell.text = "Indirect Overriding 2"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['Indirect Overriding 2'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
        
    #set Financing if exists
    if(not(pd.isnull(dfs.iloc[i]['Financing']))==True and (dfs.iloc[i]['Financing']!='0')):
        cell = table.cell(j,0)
        cell.text = "Financing"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['Financing'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
        
    #set Other if exists
    if(not(pd.isnull(dfs.iloc[i]['Others']))==True and (dfs.iloc[i]['Others']!='0')):
        cell = table.cell(j,0)
        cell.text = "Others"
        
        cell = table.cell(j,4)
        cell.text = str(dfs.iloc[i]['Others'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        j+=1
        
        #set values for others details table
        ag_code = dfs.iloc[i]['Agent Code']
        ag_name = dfs.iloc[i]['Agent Name']
        
        #create another data frame that include rows with certain agent code and name
        #df_ag = dfo.loc[dfo['Agent Code'] == ag_code]
        df_ag = dfo.loc[(dfo['Agent Code'] == ag_code) & (dfo['Agent Name']== ag_name)]
        
        #loop through the others table
        #Grab the others table
        table = doc.tables[1]
        k = 0
        sum = 0
        while(k<len(df_ag)):
            #set contest name
            cell = table.cell(k+1,0)
            cell.text = df_ag.iloc[k]['Contest Name']
            
            #set amount
            cell = table.cell(k+1,1)
            cell.text = str(df_ag.iloc[k]['Amount'])
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            sum += int(df_ag.iloc[k]['Amount'].replace(",",""))
            k+=1    
        #set total amount for others table
        cell = table.cell(17,1)
        cell.text = "{:,}".format(sum) #add commas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
    #Grab first table 
    table = doc.tables[0]

    #Set Total if exists
    if(not(pd.isnull(dfs.iloc[i]['Total']))==True and (dfs.iloc[i]['Total']!='0')):
        cell = table.cell(19,4)
        cell.text = str(dfs.iloc[i]['Total'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    #save the output docx

    if(dfs.iloc[i]['Position']=='Pioneer leader'):
        docx_path = "Compensation documents/" +"AIA_Myanmar_Agency_Compensation_PioneerLeader_Mar_"+ str(dfs.iloc[i]['Agent Code']) +".docx"
    else:
        docx_path = "Compensation documents/" +"AIA_Myanmar_Agency_Compensation_Mar_"+ str(dfs.iloc[i]['Agent Code']) +".docx"

    doc.save(docx_path)

    #time.sleep(5)

    #convert output docx to pdf
    command = ['libreoffice7.5','--convert-to','pdf','--outdir','Compensation documents',docx_path]

    subprocess.run(command)
    os.remove(docx_path)

#append path for libreoffice
libreoffice_path = "/opt/libreoffice7.5/program"
sys.path.append(libreoffice_path)

#read docx template
input_docx = 'DocxTemplate/input.docx'
doc = Document(input_docx)
#Grab first table
table = doc.tables[0]



app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # Get the file from the POST request
        file = request.files['file']
        
        # Save the file to a desired location
        filename = file.filename

        # Here, we save it to the current directory
        file.save(os.path.join("uploads",filename))
        return redirect('/progress-page')  # Redirect to progress page
    
    return render_template('index.html')

@app.route('/progress-page')
def progress_page():
    return render_template('progress.html')

@app.route('/progress')
def show_progress():
    def generate_progress():
        #Folder path for uploaded excel file
        folder_path = 'uploads'
        files = os.listdir(folder_path)

        file = files[0]
        print(file)
        file = folder_path+"/"+file
        #read excel file for Compensation (AG) table
        dfs = pd.read_excel(file,sheet_name = 'Compensation (AG)')
        
        #Data Cleaning
        #Change Headers
        dfs=dfs.rename(columns=dfs.iloc[0]).drop(dfs.index[0])
        dfs = dfs.iloc[0:]

        #Drop duplicate 'Total Column'
        dfs = dfs.iloc[:,list(range(0,26))+list(range(27,len(dfs.columns)))]

        #Change invalid values to 'Nan'
        dfs['FYC'] = dfs['FYC'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['RYC'] = dfs['RYC'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Direct Overriding'] = dfs['Direct Overriding'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Indirect Overriding 1'] = dfs['Indirect Overriding 1'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Indirect Overriding 2'] = dfs['Indirect Overriding 2'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Financing'] = dfs['Financing'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Others'] = dfs['Others'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')
        dfs['Total'] = dfs['Total'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')

        #Replace NaN values with '0'
        dfs = dfs.replace('NaN','0')

        #read excel file for Others table
        dfo = pd.read_excel(file,sheet_name = 'Others')

        dfo['Amount'] = dfo['Amount'].apply(lambda x: f"{int(math.ceil(x)):,}" if not math.isnan(x) else 'NaN')


        num_bars = len(dfs)-1  # Number of progress bars to display
        total_progress = len(dfs-1)  # Total progress value
        
        progress = 0
        while progress <= total_progress:
            yield f"data: {json.dumps({'progress': progress, 'percent': int((progress / total_progress) * 100)})}\n\n"
            createPDF(progress)
            progress += 1

            time.sleep(0.01)

        directory_name = "Compensation documents"
        archive_name = "documents2.zip"
        subprocess.run(["zip", "-r", archive_name, directory_name])
            
    return Response(generate_progress(), mimetype='text/event-stream')

@app.route('/download')
def download():
    return render_template('download.html')

@app.route('/download-file')
def downloadFile():
    return send_file('documents2.zip', as_attachment=True)


if __name__ == '__main__':
    app.run(host='',port='8082',debug=True)


